import PyPDF2
import docx
import os
from typing import Optional, Dict
from pathlib import Path

# classe pour extraire le texte des différents types de fichier 
class TextExtractor:
    @staticmethod
    def extract_from_pdf(file_path: str) -> Dict[str, any]:
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            return {
                'success': True,
                'text': text.strip(),
                'pages': len(pdf_reader.pages),
                'file_type': 'pdf'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Erreur lors de l'extraction PDF: {str(e)}",
                'file_type': 'pdf'
            }
    
    @staticmethod
    def extract_from_docx(file_path: str) -> Dict[str, any]:
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return {
                'success': True,
                'text': text.strip(),
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'file_type': 'docx'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f"Erreur lors de l'extraction DOCX: {str(e)}",
                'file_type': 'docx'
            }
    
    @staticmethod
    def extract_from_txt(file_path: str) -> Dict[str, any]:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return {
                'success': True,
                'text': text.strip(),
                'file_type': 'txt'
            }
            
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                
                return {
                    'success': True,
                    'text': text.strip(),
                    'file_type': 'txt',
                    'encoding': 'latin-1'
                }
            except Exception as e:
                return {
                    'success': False,
                    'error': f"Erreur d'encodage: {str(e)}",
                    'file_type': 'txt'
                }
        
        except Exception as e:
            return {
                'success': False,
                'error': f"Erreur lors de la lecture du fichier: {str(e)}",
                'file_type': 'txt'
            }
    
    @classmethod
    def extract_text_from_file(cls, file_path: str) -> Dict[str, any]:
        if not os.path.exists(file_path):
            return {
                'success': False,
                'error': 'Fichier non trouvé',
                'file_path': file_path
            }
        
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return cls.extract_from_pdf(file_path)
        elif file_extension == '.docx':
            return cls.extract_from_docx(file_path)
        elif file_extension in ['.txt', '.text']:
            return cls.extract_from_txt(file_path)
        else:
            return {
                'success': False,
                'error': f'Format de fichier non supporté: {file_extension}',
                'file_path': file_path
            }
    
    @staticmethod
    def clean_extracted_text(text: str) -> str:
        """
        Nettoie le texte extrait (supprime les caractères indésirables, etc.)
        """
        if not text:
            return ""
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                cleaned_lines.append(line)
        
        cleaned_text = ' '.join(cleaned_lines)
        
        import re
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        
        return cleaned_text.strip()
    
    @classmethod
    def extract_and_clean(cls, file_path: str) -> Dict[str, any]:
        """
        Extrait et nettoie le texte d'un fichier
        """
        extraction_result = cls.extract_text_from_file(file_path)
        
        if extraction_result['success']:
            extraction_result['cleaned_text'] = cls.clean_extracted_text(
                extraction_result['text']
            )
        
        return extraction_result

if __name__ == "__main__":
    extractor = TextExtractor()
    
    test_file = "test_cv.txt"
    with open(test_file, 'w', encoding='utf-8') as f:
        f.write("""
        John Doe
        Software Engineer
        
        Experience:
        - 5 years in Python development
        - Django framework expert
        - React frontend development
        
        Education:
        - Bachelor's degree in Computer Science
        - Master's in Software Engineering
        
        Skills:
        Python, Django, React, JavaScript, SQL
        """)
    
    result = extractor.extract_and_clean(test_file)
    print("Résultat de l'extraction:")
    print(result)
    
    os.remove(test_file)
